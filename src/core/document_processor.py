"""
Document processor for handling multiple file formats.
Supports DOCX, PDF, CSV, and MD files for proposals and solicitation documents.
"""

import logging
import csv
from pathlib import Path
from typing import List, Dict, Any
import docx
from marker.converters.pdf import PdfConverter
from marker.models import create_model_dict
from marker.output import text_from_rendered


class DocumentProcessor:
    """Handles processing of multiple document formats."""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def _get_processed_document_path(self, original_path: Path, doc_type: str) -> Path:
        """Get the path for a processed document."""
        if doc_type == "proposal":
            processed_dir = original_path.parent / "processed"
        else:  # solicitation
            processed_dir = original_path.parent / "processed"
        
        base_name = original_path.stem
        if doc_type == "supporting":
            return processed_dir / f"supporting_{base_name}_processed.json"
        elif doc_type == "solicitation":
            return processed_dir / f"solicitation_{base_name}_processed.json"
        else:
            return processed_dir / f"{base_name}_processed.json"
    
    def _load_processed_document(self, processed_path: Path) -> Dict[str, Any]:
        """Load a processed document from JSON."""
        try:
            import json
            with open(processed_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            self.logger.warning(f"Failed to load processed document {processed_path}: {e}")
            return None
    
    def _is_document_processed(self, original_path: Path, doc_type: str) -> bool:
        """Check if a document has already been processed."""
        processed_path = self._get_processed_document_path(original_path, doc_type)
        return processed_path.exists()
    
    def process_main_proposal(self, proposal_path: Path) -> Dict[str, Any]:
        """
        Process the main proposal file (DOCX or PDF).
        Extracts text with heading structure for DOCX.
        """
        if not proposal_path.exists():
            raise FileNotFoundError(f"Main proposal not found: {proposal_path}")
        
        # Check if already processed
        if self._is_document_processed(proposal_path, "proposal"):
            self.logger.info(f"Using cached processed document: {proposal_path}")
            processed_path = self._get_processed_document_path(proposal_path, "proposal")
            processed_data = self._load_processed_document(processed_path)
            if processed_data:
                return processed_data["content"]
        
        self.logger.info(f"Processing main proposal: {proposal_path}")
        
        if proposal_path.suffix.lower() == '.docx':
            result = self._process_docx_proposal(proposal_path)
        elif proposal_path.suffix.lower() == '.pdf':
            result = self._process_pdf_proposal(proposal_path)
        else:
            raise ValueError(f"Unsupported proposal format: {proposal_path.suffix}")
        
        # Save processed document
        self._save_processed_document(proposal_path, result, "proposal")
        
        return result
    
    def _save_processed_document(self, original_path: Path, processed_data: Dict[str, Any], doc_type: str):
        """Save processed document to processed/ subfolder."""
        try:
            # Create processed directory in the main folder (not subfolders)
            if doc_type == "proposal":
                processed_dir = original_path.parent / "processed"
            else:  # solicitation
                processed_dir = original_path.parent / "processed"
            
            processed_dir.mkdir(exist_ok=True)
            
            # Create filename for processed document
            base_name = original_path.stem
            processed_path = processed_dir / f"{base_name}_processed.json"
            
            # Save as JSON with metadata
            processed_doc = {
                "original_file": str(original_path),
                "processed_at": str(Path.cwd()),
                "format": processed_data.get("format", "unknown"),
                "content": processed_data,
                "file_size_bytes": original_path.stat().st_size if original_path.exists() else 0
            }
            
            import json
            with open(processed_path, "w", encoding="utf-8") as f:
                json.dump(processed_doc, f, indent=2, ensure_ascii=False)
            
            self.logger.info(f"Saved processed document to: {processed_path}")
            
        except Exception as e:
            self.logger.warning(f"Failed to save processed document: {e}")
    
    def _save_processed_supporting_docs(self, supporting_docs: List[Dict[str, Any]], supporting_dir: Path):
        """Save processed supporting documents as individual files."""
        try:
            processed_dir = supporting_dir.parent / "processed"
            processed_dir.mkdir(exist_ok=True)
            
            for doc in supporting_docs:
                # Create individual processed file for each supporting document
                original_path = Path(doc["file_path"])
                base_name = original_path.stem
                processed_path = processed_dir / f"supporting_{base_name}_processed.json"
                
                processed_doc = {
                    "original_file": doc["file_path"],
                    "processed_at": str(Path.cwd()),
                    "format": doc.get("format", "unknown"),
                    "content": doc,
                    "file_size_bytes": original_path.stat().st_size if original_path.exists() else 0,
                    "type": "supporting_document"
                }
                
                import json
                with open(processed_path, "w", encoding="utf-8") as f:
                    json.dump(processed_doc, f, indent=2, ensure_ascii=False)
                
                self.logger.info(f"Saved processed supporting document to: {processed_path}")
            
        except Exception as e:
            self.logger.warning(f"Failed to save processed supporting documents: {e}")
    
    def _save_processed_solicitation_docs(self, solicitation_data: Dict[str, Any], solicitation_dir: Path):
        """Save processed solicitation documents as individual files."""
        try:
            processed_dir = solicitation_dir / "processed"
            processed_dir.mkdir(exist_ok=True)
            
            for doc in solicitation_data.get("solicitation_documents", []):
                # Create individual processed file for each solicitation document
                original_path = Path(doc["file_path"])
                base_name = original_path.stem
                processed_path = processed_dir / f"solicitation_{base_name}_processed.json"
                
                processed_doc = {
                    "original_file": doc["file_path"],
                    "processed_at": str(Path.cwd()),
                    "format": doc.get("type", "unknown"),
                    "content": doc,
                    "file_size_bytes": original_path.stat().st_size if original_path.exists() else 0,
                    "type": "solicitation_document"
                }
                
                import json
                with open(processed_path, "w", encoding="utf-8") as f:
                    json.dump(processed_doc, f, indent=2, ensure_ascii=False)
                
                self.logger.info(f"Saved processed solicitation document to: {processed_path}")
            
        except Exception as e:
            self.logger.warning(f"Failed to save processed solicitation documents: {e}")
    
    def _process_docx_proposal(self, docx_path: Path) -> Dict[str, Any]:
        """Process DOCX proposal with heading structure."""
        doc = docx.Document(docx_path)
        
        # Extract text with heading structure
        sections = []
        current_section = {"title": "", "content": "", "level": 0}
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Determine heading level
            if paragraph.style.name.startswith('Heading'):
                # Save previous section if it has content
                if current_section["content"].strip():
                    sections.append(current_section.copy())
                
                # Start new section
                level = int(paragraph.style.name.replace('Heading ', ''))
                current_section = {
                    "title": text,
                    "content": text,
                    "level": level
                }
            else:
                # Add to current section content
                if current_section["content"]:
                    current_section["content"] += "\n" + text
                else:
                    current_section["content"] = text
        
        # Add final section
        if current_section["content"].strip():
            sections.append(current_section)
        
        # Combine all text for full proposal
        full_text = "\n\n".join([section["content"] for section in sections])
        
        return {
            "full_text": full_text,
            "sections": sections,
            "file_path": str(docx_path),
            "format": "docx"
        }
    
    def _process_pdf_proposal(self, pdf_path: Path) -> Dict[str, Any]:
        """Process PDF proposal."""
        converter = PdfConverter(artifact_dict=create_model_dict())
        
        # Convert PDF to markdown
        rendered = converter(str(pdf_path))
        text, _, images = text_from_rendered(rendered)
        
        return {
            "full_text": text,
            "sections": [{"title": "PDF Document", "content": text, "level": 1}],
            "file_path": str(pdf_path),
            "format": "pdf"
        }
    
    def process_supporting_docs(self, supporting_dir: Path) -> List[Dict[str, Any]]:
        """
        Process supporting documents (PDF and DOCX).
        Converts PDFs to markdown and extracts text from DOCX.
        """
        if not supporting_dir.exists():
            self.logger.warning(f"Supporting docs directory not found: {supporting_dir}")
            return []

        # Find all PDF and DOCX files (including sub-folders)
        pdf_files = list(supporting_dir.rglob("*.pdf"))
        docx_files = list(supporting_dir.rglob("*.docx"))

        all_files = pdf_files + docx_files

        if not all_files:
            self.logger.warning(f"No supporting documents found in {supporting_dir}")
            return []

        self.logger.info(f"Processing {len(all_files)} supporting documents ({len(pdf_files)} PDF, {len(docx_files)} DOCX)")

        supporting_docs = []
        needs_processing = False

        for file_path in all_files:
            # Check if already processed
            if self._is_document_processed(file_path, "supporting"):
                self.logger.info(f"Using cached processed document: {file_path.name}")
                processed_path = self._get_processed_document_path(file_path, "supporting")
                processed_data = self._load_processed_document(processed_path)
                if processed_data:
                    supporting_docs.append(processed_data["content"])
                    continue
            
            needs_processing = True
            try:
                self.logger.info(f"Processing supporting document: {file_path.name}")

                if file_path.suffix.lower() == '.pdf':
                    doc_content = self._process_pdf_document(file_path)
                elif file_path.suffix.lower() == '.docx':
                    doc_content = self._process_docx_document(file_path)
                else:
                    self.logger.warning(f"Skipping unsupported file: {file_path}")
                    continue

                doc_data = {
                    "file_path": str(file_path),
                    "file_name": file_path.name,
                    "content": doc_content,
                    "type": "supporting_document",
                    "format": file_path.suffix.lower()
                }

                supporting_docs.append(doc_data)

            except Exception as e:
                self.logger.error(f"Failed to process {file_path}: {e}")
                continue

        # Save processed supporting documents only if we processed new ones
        if needs_processing:
            self._save_processed_supporting_docs(supporting_docs, supporting_dir)

        return supporting_docs
    
    def _process_pdf_document(self, pdf_path: Path) -> str:
        """Process PDF document to text."""
        converter = PdfConverter(artifact_dict=create_model_dict())
        rendered = converter(str(pdf_path))
        text, _, images = text_from_rendered(rendered)
        return text
    
    def _process_docx_document(self, docx_path: Path) -> str:
        """Process DOCX document to text."""
        doc = docx.Document(docx_path)
        paragraphs = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)
        
        return "\n\n".join(paragraphs)
    
    def process_solicitation_docs(self, solicitation_dir: Path) -> Dict[str, Any]:
        """
        Process solicitation documents in various formats (CSV, MD, PDF).
        """
        if not solicitation_dir.exists():
            raise FileNotFoundError(f"Solicitation directory not found: {solicitation_dir}")

        self.logger.info(f"Processing solicitation documents from: {solicitation_dir}")

        # Import file discovery to find documents
        from .file_discovery import FileDiscovery
        file_discovery = FileDiscovery()
        solicitation_files = file_discovery.find_solicitation_docs(solicitation_dir)

        # Process each file type
        csv_docs = []
        md_docs = []
        pdf_docs = []
        needs_processing = False

        # Process CSV files
        for csv_path in solicitation_files["csv"]:
            # Check if already processed
            if self._is_document_processed(csv_path, "solicitation"):
                self.logger.info(f"Using cached processed document: {csv_path.name}")
                processed_path = self._get_processed_document_path(csv_path, "solicitation")
                processed_data = self._load_processed_document(processed_path)
                if processed_data:
                    csv_docs.append(processed_data["content"])
                    continue
            
            needs_processing = True
            try:
                self.logger.info(f"Processing solicitation CSV: {csv_path.name}")
                csv_content = self._process_csv_document(csv_path)
                csv_docs.append({
                    "file_path": str(csv_path),
                    "file_name": csv_path.name,
                    "content": csv_content,
                    "type": "solicitation_csv"
                })
            except Exception as e:
                self.logger.error(f"Failed to process CSV {csv_path}: {e}")
                continue

        # Process MD files
        for md_path in solicitation_files["md"]:
            # Check if already processed
            if self._is_document_processed(md_path, "solicitation"):
                self.logger.info(f"Using cached processed document: {md_path.name}")
                processed_path = self._get_processed_document_path(md_path, "solicitation")
                processed_data = self._load_processed_document(processed_path)
                if processed_data:
                    md_docs.append(processed_data["content"])
                    continue
            
            needs_processing = True
            try:
                self.logger.info(f"Processing solicitation MD: {md_path.name}")
                with open(md_path, "r", encoding="utf-8") as f:
                    md_content = f.read().strip()

                md_docs.append({
                    "file_path": str(md_path),
                    "file_name": md_path.name,
                    "content": md_content,
                    "type": "solicitation_md"
                })
            except Exception as e:
                self.logger.error(f"Failed to process MD {md_path}: {e}")
                continue

        # Process PDF files
        converter = PdfConverter(artifact_dict=create_model_dict())
        for pdf_path in solicitation_files["pdf"]:
            # Check if already processed
            if self._is_document_processed(pdf_path, "solicitation"):
                self.logger.info(f"Using cached processed document: {pdf_path.name}")
                processed_path = self._get_processed_document_path(pdf_path, "solicitation")
                processed_data = self._load_processed_document(processed_path)
                if processed_data:
                    pdf_docs.append(processed_data["content"])
                    continue
            
            needs_processing = True
            try:
                self.logger.info(f"Processing solicitation PDF: {pdf_path.name}")

                rendered = converter(str(pdf_path))
                text, _, images = text_from_rendered(rendered)

                pdf_docs.append({
                    "file_path": str(pdf_path),
                    "file_name": pdf_path.name,
                    "content": text,
                    "type": "solicitation_pdf"
                })

            except Exception as e:
                self.logger.error(f"Failed to process solicitation PDF {pdf_path}: {e}")
                continue

        # Combine all solicitation documents
        all_solicitation_docs = csv_docs + md_docs + pdf_docs

        # Save processed solicitation documents only if we processed new ones
        if needs_processing:
            self._save_processed_solicitation_docs({"solicitation_documents": all_solicitation_docs, "csv_documents": csv_docs, "md_documents": md_docs, "pdf_documents": pdf_docs}, solicitation_dir)

        return {
            "solicitation_documents": all_solicitation_docs,
            "csv_documents": csv_docs,
            "md_documents": md_docs,
            "pdf_documents": pdf_docs
        }
    
    def _process_csv_document(self, csv_path: Path) -> str:
        """Process CSV document to text format."""
        try:
            with open(csv_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = list(reader)
            
            # Convert CSV to readable text format
            text_lines = []
            for i, row in enumerate(rows):
                if i == 0:  # Header row
                    text_lines.append(" | ".join(row))
                    text_lines.append("-" * len(" | ".join(row)))
                else:
                    text_lines.append(" | ".join(row))
            
            return "\n".join(text_lines)
        except Exception as e:
            self.logger.error(f"Failed to process CSV {csv_path}: {e}")
            return f"Error processing CSV file: {csv_path.name}" 